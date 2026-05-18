import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from io import BytesIO
import time
import re

# --- 1. KONFIGURASI HALAMAN WEB STREAMLIT ---
st.set_page_config(page_title="Sistem Perangkat Pembelajaran KBC MAN 2 Makassar", layout="wide")

# --- 2. SETUP API KEY & NAVIGASI DI SIDEBAR ---
st.sidebar.title("🔑 Pengaturan & Navigasi")
api_key = st.sidebar.text_input("Masukkan Google Gemini API Key:", type="password")

if api_key:
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel('gemini-2.5-flash')
else:
    st.sidebar.warning("Silakan masukkan API Key di sini untuk mengaktifkan AI.")

st.sidebar.markdown("---")
st.sidebar.subheader("📌 Pilih Modul Kerja")

# Navigasi Utama Aplikasi (Tiga menu mandiri)
menu_utama = st.sidebar.radio(
    "Silakan pilih layanan:",
    ["📝 Menyusun RPP", "📊 Menyusun Kisi-Kisi", "❓ Membuat Soal"]
)

# --- 3. JUDUL UTAMA APLIKASI ---
st.title("🍎 Sistem Perencanaan & Evaluasi Pembelajaran (KBC)")
st.caption("Edisi Eco-Friendly (Pembelajaran Mendalam): MAN 2 Kota Makassar")

# --- 4. INISIALISASI SESSION STATE DATA ---
if 'rpp_output' not in st.session_state:
    st.session_state.rpp_output = ""
if 'kisi_output' not in st.session_state:
    st.session_state.kisi_output = ""
if 'soal_output' not in st.session_state:
    st.session_state.soal_output = ""


# --- 5. FUNGSI CORE AI ---
def panggil_ai(prompt):
    if not api_key:
        st.error("Masukkan API Key terlebih dahulu di sidebar kiri!")
        return ""
    try:
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        if "429" in str(e) or "quota" in str(e).lower():
            st.warning("⚠️ Kuota menit ini habis. Mohon tunggu 60 detik...")
            time.sleep(5)
        else:
            st.error(f"Terjadi kesalahan: {str(e)}")
        return ""

# --- 6. FUNGSI FORMAT WORD EKSPOR (ELEGAN & ECO-FRIENDLY) ---
def set_cell_background(cell, color_hex):
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def format_cell_margins(cell, top=60, bottom=60, left=100, right=100):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def pasang_garis_pembatas_tabel(table, color_hex="CBD5E1"):
    tblPr = table._tbl.tblPr
    borders_xml = f"""
    <w:tblBorders {nsdecls("w")}>
        <w:top w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>
        <w:bottom w:val="single" w:sz="6" w:space="0" w:color="{color_hex}"/>
        <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>
        <w:left w:val="none"/>
        <w:right w:val="none"/>
        <w:insideV w:val="none"/>
    </w:tblBorders>
    """
    tblPr.append(parse_xml(borders_xml))

def buat_dokumen_word_kbc(judul_dokumen, metadata, isi_konten):
    doc = Document()
    
    # MARGIN SUPER SEMPIT (1.5 cm / 0.6 Inci) - Eco-Friendly & Tetap Elegan
    for section in doc.sections:
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    # HEADER IDENTITAS UTAMA (DESAIN ELEGAN)
    header_p = doc.add_paragraph()
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_p.paragraph_format.space_after = Pt(2)
    
    h_run = header_p.add_run(f"{judul_dokumen}\n")
    h_run.font.name = 'Georgia'
    h_run.font.size = Pt(13)
    h_run.font.bold = True
    h_run.font.color.rgb = RGBColor(26, 54, 93)  # Navy Blue
    
    sub_run = header_p.add_run("Integrasi Ekoteologi dalam Kurikulum Berbasis Cinta (KBC) — MAN 2 Kota Makassar\n")
    sub_run.font.name = 'Arial'
    sub_run.font.size = Pt(9.5)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(115, 125, 140)  # Muted Gray

    # Tabel Ringkas Identitas Madrasah (Desain Clean)
    meta_table = doc.add_table(rows=3, cols=4)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    pasang_garis_pembatas_tabel(meta_table, "CBD5E1")  # Garis abu-abu terang yang halus
    
    identitas_data = [
        ("Nama Madrasah", f": {metadata.get('madrasah', 'MAN 2 KOTA MAKASSAR')}", "Nama Guru", f": {metadata.get('guru', 'Kaharuddin, S.Pd')}"),
        ("Mata Pelajaran", f": {metadata.get('mapel', 'PPKn')}", "Fase/Kelas/Smt", f": {metadata.get('kelas', 'F / XII / Ganjil')}"),
        ("Topik Pokok", f": {metadata.get('topik', '')}", "Jenis Asesmen / Waktu", f": {metadata.get('waktu', '')}")
    ]
    
    for i, (k1, v1, k2, v2) in enumerate(identitas_data):
        row = meta_table.rows[i]
        row.cells[0].text = k1
        row.cells[1].text = v1
        row.cells[2].text = k2
        row.cells[3].text = v2
        for cell in row.cells:
            format_cell_margins(cell, top=50, bottom=50, left=90, right=90)
            for p in cell.paragraphs:
                p.paragraph_format.line_spacing = 1.0
                for r in p.runs:
                    r.font.name = 'Arial'
                    r.font.size = Pt(9.5)
                    r.font.color.rgb = RGBColor(71, 85, 105)  # Slate Gray

    # Garis pembatas horizontal hiasan di bawah header
    border_p = doc.add_paragraph()
    border_p.paragraph_format.space_before = Pt(4)
    border_p.paragraph_format.space_after = Pt(8)
    pBdr = OxmlElement('w:pBdr')
    bottom_border = OxmlElement('w:bottom')
    bottom_border.set(qn('w:val'), 'single')
    bottom_border.set(qn('w:sz'), '6')
    bottom_border.set(qn('w:space'), '1')
    bottom_border.set(qn('w:color'), '1A365D') # Garis Navy Blue tipis
    pBdr.append(bottom_border)
    border_p._p.get_or_add_pPr().append(pBdr)

    # PARSING KONTEN TEKS & TABEL MARKDOWN
    baris_list = isi_konten.split('\n')
    di_dalam_tabel = False
    tabel_obj = None
    
    for baris in baris_list:
        baris_bersih = baris.strip()
        if not baris_bersih:
            continue
        
        # Logika Konversi Tabel otomatis
        if baris_bersih.startswith('|'):
            kolom_data = [k.strip() for k in baris_bersih.split('|')[1:-1]]
            if not kolom_data or all(c == '' or c.startswith('-') for c in kolom_data):
                continue 
            
            if not di_dalam_tabel:
                di_dalam_tabel = True
                tabel_obj = doc.add_table(rows=0, cols=len(kolom_data))
                tabel_obj.alignment = WD_TABLE_ALIGNMENT.CENTER
                pasang_garis_pembatas_tabel(tabel_obj, "CBD5E1")
                
                # Desain Header Tabel (Biru Dongker, Teks Putih Bold)
                row = tabel_obj.add_row()
                for idx, teks in enumerate(kolom_data):
                    cell = row.cells[idx]
                    cell.text = teks
                    set_cell_background(cell, "1A365D")  # Deep Navy
                    format_cell_margins(cell, top=80, bottom=80, left=100, right=100)
                    for p_cell in cell.paragraphs:
                        p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        if p_cell.runs:
                            p_cell.runs[0].font.bold = True
                            p_cell.runs[0].font.color.rgb = RGBColor(255, 255, 255)
                            p_cell.runs[0].font.name = 'Arial'
                            p_cell.runs[0].font.size = Pt(9.5)
                continue
            
            # Isian Baris Tabel Konten
            row = tabel_obj.add_row()
            for idx, teks in enumerate(kolom_data):
                cell = row.cells[idx]
                cell.text = teks
                format_cell_margins(cell, top=60, bottom=60, left=90, right=90)
                for p_cell in cell.paragraphs:
                    p_cell.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for r in p_cell.runs:
                        r.font.name = 'Arial'
                        r.font.size = Pt(9.5)
                        r.font.color.rgb = RGBColor(51, 65, 85)
        else:
            di_dalam_tabel = False
            
            # Desain Judul Bab / Komponen Utama
            if baris_bersih.startswith(('A.', 'B.', 'C.', 'D.', 'E.', 'F.', 'G.', 'H.', 'I.', 'J.', 'K.', 'Langkah', 'AWAL', 'INTI', 'PENUTUP', 'ALAT', '1.', '2.', '3.')):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.keep_with_next = True
                
                run = p.add_run(baris_bersih)
                run.font.bold = True
                run.font.name = 'Georgia'
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(26, 54, 93)  # Navy Blue
            else:
                # Format Paragraf Isi Teks Biasa (Justify)
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.line_spacing = 1.15
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                parts = re.split(r'(\*\*.*?\*\*)', baris_bersih)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(15, 23, 42)
                    else:
                        run = p.add_run(part)
                        run.font.color.rgb = RGBColor(51, 65, 85)
                    run.font.name = 'Arial'
                    run.font.size = Pt(10)
                    
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


# =========================================================================
# 📝 MENU 1: MENYUSUN RPP
# =========================================================================
if menu_utama == "📝 Menyusun RPP":
    st.header("Modul Perencanaan Pembelajaran Mendalam & KBC")
    st.subheader("📋 Formulir Input Data Utama RPP")
    
    col1, col2 = st.columns(2)
    with col1:
        v_madrasah = st.text_input("Nama Madrasah:", value="MAN 2 KOTA MAKASSAR")
        v_guru = st.text_input("Nama Guru:", value="Kaharuddin, S.Pd")
        v_mapel = st.text_input("Mata Pelajaran:", value="PPKn")
        v_kelas = st.text_input("Fase / Kelas / Semester:", value="F / XII / Ganjil")
    with col2:
        v_topik = st.text_input("Topik Pembelajaran:", placeholder="Contoh: Pancasila dalam Kehidupan Global")
        v_subtopik = st.text_input("Sub Topik Pembelajaran:", placeholder="Contoh: Tantangan Ideologi Pancasila di Era Digital")
        v_tujuan = st.text_area("Tujuan Pembelajaran Umum:", placeholder="Contoh: Menganalisis peluang dan tantangan penerapan nilai-nilai Pancasila dalam kehidupan global...")
        v_waktu = st.text_input("Alokasi Waktu / Jumlah Pertemuan:", value="2 x 45 Menit / 1 Pertemuan")

    st.markdown("---")
    
    st.subheader("🖥️ Hasil Penyusunan Dokumen RPP")
    st.session_state.rpp_output = st.text_area(
        "Draf Teks RPP (Dapat diedit manual sebelum diunduh):", 
        value=st.session_state.rpp_output, 
        height=400
    )
    
    if st.button("✨ Generate RPP Sesuai Aturan KBC & Deep Learning"):
        if not v_topik or not v_subtopik or not v_tujuan:
            st.error("Silakan lengkapi Topik, Sub-Topik, dan Tujuan Pembelajaran terlebih dahulu!")
        else:
            with st.spinner("AI sedang merumuskan draf RPP terstruktur, padat, dan hemat kertas..."):
                prompt_rpp = f"""
                Saya seorang guru PPKN jenjang SMA/MA. Buatkan saya RPP sesuai sub materi dan tujuan pembelajaran berikut:
                Nama Madrasah: {v_madrasah}
                Nama Guru: {v_guru}
                Mapel: {v_mapel}
                Fase/Kelas/Smt: {v_kelas}
                Topik Pembelajaran: {v_topik}
                Sub Topik Pembelajaran: {v_subtopik}
                Tujuan pembelajaran umum: {v_tujuan}
                Alokasi Waktu /pertemuan: {v_waktu}

                Buat bagian awal rencana pembelajaran dengan spesifikasi berikut secara padat, langsung inti, tanpa spasi berlebih demi hemat kertas:
                A. Identifikasi Peserta Didik (Kesiapan, pengetahuan awal, minat, latar belakang, kebutuhan. Tulis bentuk poin).
                B. Identifikasi Materi Pembelajaran (Jenis pengetahuan, relevansi kehidupan nyata, tingkat kesulitan, struktur materi, integrasi nilai/karakter. Tulis bentuk poin).
                C. Dimensi Profil Lulusan (Pilih 4 dimensi di antara: Keimanan dan Ketakwaan, Kewargaan, Penalaran Kritis, Kreativitas, Kolaborasi, Kemandirian, Kesehatan, Komunikasi).
                D. Topik Panca Cinta (Pilih 3 topik di antara: Cinta Allah dan Rasul-Nya, Cinta Ilmu, Cinta Lingkungan, Cinta Diri dan Sesama Manusia, Cinta Tanah Air).
                E. Materi Integrasi KBC (Tuliskan materi integrasi ekoteologi dalam Kurikulum Berbasis Cinta (KBC) yang relevan dalam bentuk paragraf).
                F. Topik Pembelajaran (Relevan dengan capaian dan tujuan).
                G. Tujuan Pembelajaran (Tulis tujuan umum, lalu tujuan rinci per pertemuan mencakup subjek, pengetahuan, keterampilan/sikap dengan KKO terukur).
                H. Praktik Pedagogis (Tentukan Model/Strategi/Metode yang sesuai dengan sub topik untuk mencapai dimensi profil lulusan).
                I. Kemitraan Pembelajaran (Mitra kolaborasi di lingkungan sekolah, luar sekolah, atau masyarakat).
                J. Lingkungan Pembelajaran (Integrasi ruang fisik, virtual, dan budaya belajar).
                K. Pemanfaatan Digital (Pemanfaatan teknologi digital interaktif dan kontekstual).

                Berdasarkan hasil di atas, buat langkah pembelajaran yang memuat prinsip pembelajaran mendalam (berkesadaran, bermakna, menggembirakan) sesuai jumlah pertemuan di atas:
                Langkah – Langkah Pembelajaran:
                - AWAL (Sebutkan prinsip pembelajaran. Pembuka, salam, cek kehadiran, kondisi siap belajar, orientasi bermakna, apersepsi kontekstual, penyampaian tujuan, motivasi menggembirakan).
                - INTI:
                  1. Memahami (Sebutkan prinsip pembelajaran. Membangun kesadaran murid terhadap tujuan, aktif mengonstruksi pengetahuan mendalam konsep/materi dari berbagai sumber).
                  2. Mengaplikasi (Sebutkan prinsip pembelajaran. Menerapkan pengetahuan individu/kolaboratif melalui pemecahan masalah, pengambilan keputusan, dll).
                  3. Merefleksi (Sebutkan prinsip pembelajaran. Murid mengevaluasi dan memaknai proses/hasil praktik nyata, sejauh mana tujuan tercapai, kekuatan, tantangan, area perbaikan).
                - PENUTUP (Sebutkan prinsip pembelajaran. Umpan balik konstruktif, menyimpulkan pembelajaran, keterlibatan perencanaan pembelajaran selanjutnya).

                ALAT DAN MEDIA PEMBELAJARAN
                (Tentukan Alat dan Media Pembelajaran sesuai dengan sub topik dan Tujuan pembelajaran setiap pertemuan).
                """
                hasil = panggil_ai(prompt_rpp)
                if hasil:
                    st.session_state.rpp_output = hasil
                    st.success("RPP berhasil dibuat! Silakan cek kolom teks di atas.")
                    st.rerun()

    st.markdown("---")
    if st.session_state.rpp_output:
        meta = {"madrasah": v_madrasah, "guru": v_guru, "mapel": v_mapel, "kelas": v_kelas, "topik": v_topik, "waktu": v_waktu}
        file_doc = buat_dokumen_word_kbc("PERENCANAAN PEMBELAJARAN MENDALAM (RPP / MODUL AJAR)", meta, st.session_state.rpp_output)
        st.download_button(
            label="📥 Download RPP (.docx) - Eco-Friendly", 
            data=file_doc, 
            file_name=f"RPP_KBC_{v_subtopik.replace(' ', '_')}.docx"
        )


# =========================================================================
# 📊 MENU 2: MENYUSUN KISI-KISI SOAL
# =========================================================================
elif menu_utama == "📊 Menyusun Kisi-Kisi":
    st.header("Modul Penyusunan Kisi-Kisi Instrumen Soal")
    st.subheader("📋 Formulir Input Data Kisi-Kisi")
    
    col1, col2 = st.columns(2)
    with col1:
        k_madrasah = st.text_input("Nama Madrasah:", value="MAN 2 KOTA MAKASSAR", key="k_m")
        k_guru = st.text_input("Nama Guru:", value="Kaharuddin, S.Pd", key="k_g")
        k_mapel = st.text_input("Mata Pelajaran:", value="PPKn", key="k_mp")
    with col2:
        k_kelas = st.text_input("Kelas / Semester:", value="XII / Ganjil", key="k_kl")
        k_materi = st.text_input("Materi Pokok / Sub Materi:", placeholder="Contoh: Tantangan Global Penegakan HAM di Indonesia")
        k_tujuan = st.text_area("Tujuan Pembelajaran Soal Ujian:", placeholder="Contoh: Peserta didik mampu menganalisis dilema penegakan hukum dalam menjaga harmoni sosial...")

    st.markdown("---")
    
    st.subheader("🖥️ Hasil Penyusunan Matriks Kisi-Kisi")
    st.session_state.kisi_output = st.text_area(
        "Draf Teks Kisi-Kisi (Format Tabel):", 
        value=st.session_state.kisi_output, 
        height=400
    )
    
    if st.button("✨ Generate Matriks Kisi-Kisi (Deep Learning + KBC)"):
        if not k_materi or not k_tujuan:
            st.error("Silakan lengkapi Materi Pokok dan Tujuan Pembelajaran!")
        else:
            with st.spinner("AI sedang menyusun matriks tabel kisi-kisi HOTS & KBC..."):
                prompt_kisi = f"""
                Bertindaklah sebagai seorang Ahli Kurikulum, Desainer Instruksional, dan Pengembang Evaluasi Pembelajaran. 
                Susun Kisi-Kisi Penulisan Soal yang mengintegrasikan Pendekatan Deep Learning (Mindful, Meaningful, Joyful Learning) dan esensi Kurikulum Berbasis Cinta (KBC).

                Data awal:
                - Mata Pelajaran: {k_mapel}
                - Kelas / Semester: {k_kelas}
                - Materi Pokok : {k_materi}
                - Tujuan Pembelajaran : {k_tujuan}

                Ketentuan Penyusunan Kisi-Kisi:
                1. Indikator Soal (IKTP): Diarahkan pada kemampuan berpikir kritis (Mindful), pemecahan masalah kontekstual nyata sehari-hari (Meaningful), memicu rasa ingin tahu tanpa rasa takut (Joyful).
                2. Dimensi KBC: Sisipkan aspek Kurikulum Berbasis Cinta seperti empati, kepedulian sosial/lingkungan, gotong royong, kasih sayang, atau nilai kemanusiaan dalam stimulus soal.
                3. Bentuk Stimulus: Harus berbasis kasus nyata, cerita inspiratif, dilema moral, atau fenomena sehari-hari yang dekat dengan dunia anak.

                Format Output Kisi-Kisi wajib berupa TABEL MARKDOWN dengan kolom-kolom berikut:
                1. No. | 2. Tujuan Pembelajaran (TP) | 3. Materi Pokok | 4. Indikator Soal | 5. Dimensi Deep Learning & KBC yang Tersirat | 6. Level Kognitif (C4/C5/C6) | 7. Bentuk Soal | 8. No. Soal

                Berikan juga 1 contoh konkret draf soal berdasarkan salah satu baris kisi-kisi yang Anda buat agar saya mendapatkan gambaran utuh. Gunakan bahasa yang hangat, profesional, dan inspiratif. Buat sepadat mungkin demi hemat kertas cetak.
                """
                hasil = panggil_ai(prompt_kisi)
                if hasil:
                    st.session_state.kisi_output = hasil
                    st.success("Kisi-kisi berhasil dibuat!")
                    st.rerun()

    st.markdown("---")
    if st.session_state.kisi_output:
        meta = {"madrasah": k_madrasah, "guru": k_guru, "mapel": k_mapel, "kelas": k_kelas, "topik": k_materi, "waktu": "-"}
        file_doc = buat_dokumen_word_kbc("KISI-KISI PENULISAN SOAL EVALUASI", meta, st.session_state.kisi_output)
        st.download_button(
            label="📥 Download Kisi-Kisi (.docx) - Eco-Friendly", 
            data=file_doc, 
            file_name=f"Kisi_Kisi_KBC_{k_materi.replace(' ', '_')}.docx"
        )


# =========================================================================
# ❓ MENU 3: MEMBUAT SOAL EVALUASI
# =========================================================================
elif menu_utama == "❓ Membuat Soal":
    st.header("Modul Pembuatan Lembar Soal & Kunci Jawaban")
    st.subheader("📋 Formulir Input Spesifikasi Bank Soal")
    
    col1, col2 = st.columns(2)
    with col1:
        s_madrasah = st.text_input("Nama Madrasah:", value="MAN 2 KOTA MAKASSAR", key="s_m")
        s_guru = st.text_input("Nama Guru:", value="Kaharuddin, S.Pd", key="s_g")
        s_mapel = st.text_input("Mata Pelajaran:", value="PPKn", key="s_mp")
        s_kelas = st.text_input("Kelas / Semester:", value="XII / Ganjil", key="s_kl")
        
        s_jenis_asesmen = st.selectbox(
            "Jenis Asesmen:", 
            ["Asesmen Formatif", "Asesmen Sumatif", "Penilaian Harian", "Ujian Tengah Semester (UTS)", "Ujian Akhir Semester (UAS)"],
            key="s_ja"
        )
    with col2:
        s_materi = st.text_input("Materi Pokok / Bahasan Ujian:", placeholder="Contoh: Implementasi Etika Konstitusi dalam Kehidupan Berbangsa")
        s_jumlah_pg = st.text_input("Jumlah Soal Pilihan Ganda Kompleks:", value="3")
        s_jumlah_bs = st.text_input("Jumlah Soal Tipe Benar / Salah:", value="2", key="s_bs")
        s_jumlah_esai = st.text_input("Jumlah Soal Esai / Studi Kasus HOTS:", value="2")

    st.markdown("---")
    
    st.subheader("🖥️ Hasil Bank Soal Evaluasi KBC")
    st.session_state.soal_output = st.text_area(
        "Naskah Lembar Soal & Kunci Jawaban:", 
        value=st.session_state.soal_output, 
        height=400
    )
    
    if st.button("✨ Generate Lembar Soal Mengetuk Hati (HOTS + KBC)"):
        if not s_materi:
            st.error("Silakan lengkapi kolom Materi Pokok terlebih dahulu!")
        else:
            with st.spinner("AI sedang merancang paket soal narasi-numerasi humanis mendalam..."):
                prompt_soal = f"""
                Bertindaklah sebagai pembuat soal evaluasi profesional. Buatlah paket soal untuk:
                - Nama Madrasah: {s_madrasah}
                - Nama Guru: {s_guru}
                - Mata Pelajaran: {s_mapel}
                - Kelas / Semester: {s_kelas}
                - Jenis Asesmen: {s_jenis_asesmen}
                - Materi Pokok: {s_materi}
                - Target Komposisi Soal: 
                  1. {s_jumlah_pg} Soal Pilihan Ganda Kompleks
                  2. {s_jumlah_bs} Soal Tipe Benar / Salah
                  3. {s_jumlah_esai} Soal Esai / Studi Kasus HOTS.

                Ketentuan dan Karakteristik Soal yang Harus Diikuti:
                1. Elemen Deep Learning:
                   - Mindful: Soal harus mengajak siswa berpikir kritis, berefleksi, dan menyadari dampak dari ilmu tersebut terhadap diri sendiri atau lingkungan.
                   - Meaningful: Soal harus berbasis studi kasus nyata (kontekstual), problem-solving, dan relevan dengan kehidupan sehari-hari siswa (bukan hafalan mati).
                   - Joyful: Soal disajikan dengan narasi yang menarik, menantang rasa ingin tahu, menggunakan bahasa yang memotivasi, dan tidak memicu kecemasan (anxiety-free assessment).

                2. Elemen Kurikulum Berbasis Cinta (KBC):
                   - Soal harus menginsersikan nilai cinta kasih (cinta sesama, cinta lingkungan, cinta tanah air, atau cinta ilmu pengetahuan).
                   - Mengandung muatan empati, kepedulian sosial, kolaborasi, atau solusi humanis.
                   - Pilihan jawaban atau rubrik penilaian harus menghargai proses berpikir dan sisi kemanusiaan siswa.
                   - Untuk opsi Pilihan Ganda Kompleks, susun pilihan jawaban A, B, C, D, E secara horizontal/menyamping untuk menghemat baris kertas.
                   - Untuk opsi Benar / Salah, sertakan ruang singkat bagi siswa untuk memberikan alasan reflektif mereka mengapa memilih jawaban tersebut agar aspek Mindful tercapai.

                Format Output yang Saya Inginkan:
                - Buatkan soal yang menyentuh hati dan penalaran tingkat tinggi (HOTS), bersifat narasi atau numerasi.
                - Urutkan nomor soal dari Pilihan Ganda Kompleks, Benar/Salah, lalu Esai.
                - Lengkapi setiap nomor soal dengan komponen wajib berikut:
                  a. Stimulus/Konteks (Cerita atau skenario yang relevan).
                  b. Pertanyaan yang diajukan.
                  c. Kunci Jawaban / Rubrik Penilaian Reflektif.
                  d. Penjelasan Mengapa Soal Ini Memenuhi Unsur Deep Learning & KBC.

                Gunakan bahasa yang santun, hangat, inspiratif, dan mudah dipahami oleh siswa. Buat dokumen sepadat mungkin demi kebaikan lingkungan (hemat kertas).
                """
                hasil = panggil_ai(prompt_soal)
                if hasil:
                    st.session_state.soal_output = hasil
                    st.success("Paket Soal Evaluasi KBC berhasil dirumuskan!")
                    st.rerun()

    st.markdown("---")
    if st.session_state.soal_output:
        meta = {"madrasah": s_madrasah, "guru": s_guru, "mapel": s_mapel, "kelas": s_kelas, "topik": s_materi, "waktu": s_jenis_asesmen}
        file_doc = buat_dokumen_word_kbc("LEMBAR EVALUASI SISWA & KUNCI JAWABAN", meta, st.session_state.soal_output)
        st.download_button(
            label="📥 Download Lembar Soal (.docx) - Eco-Friendly", 
            data=file_doc, 
            file_name=f"Soal_{s_jenis_asesmen.replace(' ', '_')}_{s_materi.replace(' ', '_')}.docx"
        )
